import re
import docx
from pyquery import PyQuery as pq


# 读取合同html页面内容
def read_contact_html(all_link):
    for html_link in all_link:
        # 从指定的url去解析合同界面
        doc = pq("https:" + html_link, encoding="gb2312")
        # 合同起始位置（该网站一个页面有两个合同范本，需要进行切割，根据页面元素找到不同范本起始位置）
        text_name = [i.text() for i in doc.items('p strong')]
        # 如果找不到合同起始位置，那说明该模板只有一个合同，所以就从另外的地方获取标题
        if len(text_name) <= 0:
            text_name = [i.text() for i in doc.items('#ArtContent h3')]
            if not text_name:
                text_name = [i.text() for i in doc.items('#ArtContent h1')]

        #     用正则匹配清洗掉小标题为代码中默认的大标题情况
        real_name = list(filter(lambda x: re.search('\.|。|、|：|:|;|；|第|条|\?', x) == None, text_name))

        # 该页面所有标签为p的文本内容内容
        text_body = [b.text() for b in doc.items('p')]
        # 需要提取的合同文本内容结束位置
        text_end = [e.text() for e in doc.items('.sfont')]

        text_head = [h.text() for h in doc.items('.ArtContent h1')]

        save_contact(real_name, text_body, text_end)


# 将读取的文件以结构化形式保存下来
def save_contact(real_name, text_body, text_end):
    range_idex = 0
    # 对解析好的合同页面内容进行切割存储
    for i in range(len(real_name)):
        file = docx.Document()
        file.add_heading(real_name[i])
        for contract_body in range(range_idex, len(text_body)):
            # 判断合同文本裁剪位置
            if (i < (len(real_name) - 1) and (text_body[contract_body] == real_name[i + 1])):
                range_idex = contract_body
                print("contract_body:" + str(contract_body) + "=" + text_body[contract_body])
                break
            #  判断文本结束位置
            if (len(text_end) >= 1 and text_body[contract_body] == text_end[0]):
                range_idex = len(text_body)
                break
            # 按段存储进docx文档中
            file.add_paragraph(text_body[contract_body])
        title_name = real_name[i] + ".docx"
        file.save(title_name)


# 通过目录页找到合同url列表
def run(url):
    # based_url = "https://www.diyifanwen.com/fanwen/zulinhetong/index_"
    # for url_index in range(11, 21):
    # url = based_url + str(url_index) + ".html"
    # 获取当前目录页面中的所有合同url
    doc_all = pq(url, encodind="gb2312", headers={'User-Agent': 'Mozilla/5.0 '
                                                                '(Windows NT 10.0; Win64; x64) AppleWebKit/537.36 '
                                                                '(KHTML, like Gecko) Chrome/46.0.2490.71 Safari/537.36'})
    all_link = [link.attr('href') for link in doc_all.items('#AListBox a')]
    read_contact_html(all_link)


# 运行
if __name__ == '__main__':
    based_url = "https://www.diyifanwen.com/fanwen/laodonghetong/"
    run(based_url)
    for url_index in range(2, 20):
        try:
            run(based_url + "index_" + str(url_index) + ".html")
        except Exception:
            continue
