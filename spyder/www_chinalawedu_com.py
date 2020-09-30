import re
import docx
from pyquery import PyQuery as pq
import os
import datetime


def c_print(content):
    if False:
        print(content)


def get_website(page_url):
    """
    :param page_url: http://www.chinalawedu.com/web/193/wa1904106573.shtml
    :return: http://www.chinalawedu.com
    """
    if page_url.index('http') == 0:
        page_url = page_url.replace('\\', '/')
        parts = re.split('/', page_url)
        parts = list(filter(lambda x: x is not None and len(x) > 0, parts))
        if len(parts) < 2:
            raise ValueError('无法解析DNS:' + str(page_url))
        return parts[0] + '//' + parts[1]


def get_data_page_urls(url, path=None):
    html = pq(url)
    all_links_elements = html('ul.info-list.clearfix li a')  # 组件名为info-list clearfix的ul中的 li组件中的a元素
    return [get_data_from_html_element(link, 'href') for link in all_links_elements]


def get_visit_data_url(page_url, data_uri):
    """
    根据不同的页面有不同的访问具体的数据页面的封装方法
    :param page_url: 'http://www.chinalawedu.com/web/193/page1.shtm'
    :param data_uri: '/web/193/wa1904106573.shtml'
    :return: http://www.chinalawedu.com/web/193/wa1904106573.shtml
    """
    dns = get_website(page_url)
    return dns + data_uri


def get_data_from_html_element(html_element, attribute_name=None):
    try:
        if attribute_name is None:
            return html_element.text
        data = html_element.attrib[attribute_name]
        return data
    except Exception as e:
        c_print(e)
        return html_element.text


def get_doc_from_html(url):
    html = pq(url)
    data_div_list = html('div#fontzoom.new-con.f14 p')  # 需要修改的地方，不同的合同数据位置不一样
    file = docx.Document()

    head_div = html('div.hd.tc h3')  # 需要修改的地方，页面，head位置不一样
    head = get_data_from_html_element(head_div[0])
    file.add_heading(head)

    for data_div in data_div_list:
        text = get_data_from_html_element(data_div)
        if isinstance(text, str):
            file.add_paragraph(text)
    return file, head


page_url = 'http://www.chinalawedu.com/web/192/page1.shtm'
def run(page_url, store_file):
    total = 0
    print('start ' + str(datetime.datetime.now()))
    for i in range(1, 114):
        page_url = page_url + str(i) + '.shtm'
        try:
            all_data_links = get_data_page_urls(page_url)
            for link in all_data_links:
                try:
                    test_url = get_visit_data_url(page_url, link)
                    docx_data, docname = get_doc_from_html(test_url)
                    docname = docname.replace("\\", '_').replace('/', '_').replace("|", '_')
                    # docx_data.save(os.path.join('G:/正保法律实务-借款', docname + '.doc'))
                    docx_data.save(os.path.join(store_file, docname + '.doc'))
                    total = total + 1
                except Exception as e:
                    print(e)
                    continue
                if total % 100 == 0:
                    print(str(datetime.datetime.now()) + 'has download ' + str(total))
        except Exception as e:
            print(e)
            continue
        print('has download page:' + str(i))

print('end ' + str(datetime.datetime.now()))

